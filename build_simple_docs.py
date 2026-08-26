import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DOC_PATH = r"d:\all_flutter_projects\cashback_reward_app\CashKaro_Project_Overview_and_Architecture.docx"
MD_PATH = r"d:\all_flutter_projects\cashback_reward_app\PROJECT_DOCUMENTATION.md"

def set_cell_background(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_padding(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = RGBColor(30, 144, 255) # CashKaro Blue
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate Navy
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 65, 85)
        run.font.bold = True
    return h

def build_word_doc():
    doc = Document()

    # Set 0.75 in margins
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Main Title
    title_p = doc.add_paragraph()
    r_title = title_p.add_run("CashKaro App — Complete Architecture, Workflow & Screens Guide")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 144, 255)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    r_sub = sub_p.add_run("A clear, concise, plain-language project manual covering architecture, folder design, app flow, and all 30 screens.\n")
    r_sub.font.size = Pt(10.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("―" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: WHAT THIS APP DOES
    add_heading(doc, "1. What This App Does (In Simple Words)", level=1)
    doc.add_paragraph(
        "CashKaro is a shopping rewards and cashback mobile app. When users want to buy anything online "
        "(clothes, electronics, medicines, groceries, or apply for personal loans/credit cards), they open "
        "CashKaro, click on a store (like Amazon, Flipkart, Myntra, or Navi Loans), and complete their purchase. "
        "The retailer pays CashKaro a commission, and CashKaro gives that money back to the user as Real Cashback. "
        "Once verified, users can transfer that money directly into their Bank Account (via NEFT), UPI (GPay/PhonePe), "
        "or get Amazon Pay gift vouchers."
    )

    doc.add_paragraph()

    # SECTION 2: WHY WE USE THIS FOLDER STRUCTURE
    add_heading(doc, "2. Why We Use This Specific Folder Structure", level=1)
    doc.add_paragraph(
        "In Flutter development, keeping all code in one place makes projects messy and difficult to fix. "
        "We divided the app into 6 distinct, organized folders so any developer can easily locate, modify, "
        "and maintain features without breaking anything else:"
    )

    folders = [
        ("lib/screens/", "Full App Pages", 
         "Every file here represents a complete screen that the user sees (e.g. Login, Home, Profile, Withdraw). "
         "Why: Keeps page-level layout, scrolling, and navigation strictly separated so you know exactly which file to open when editing a specific screen."),
        
        ("lib/widgets/", "Reusable UI Building Blocks", 
         "Contains smaller UI parts used in multiple screens (like CustomAppBar, ModalDragHandle, and ImageLoaders). "
         "Why: Instead of writing the exact same 40 lines of App Bar code on 20 different screens, we write it once here and reuse it everywhere. If we need to change the back-button color, we update it in 1 place instead of 20."),
        
        ("lib/models/", "Data Blueprints (Contracts)", 
         "Defines what data looks like in code (e.g. what a Product has: name, price, cashback, image URL). "
         "Why: Prevents bugs and spelling mistakes. When data comes from an API or database, models turn raw JSON into safe, strongly-typed Dart objects."),
        
        ("lib/providers/", "App State & Memory Management", 
         "Holds live data in memory (e.g. is user logged in, current wallet balance, active theme, selected category). "
         "Why: When a user edits their name in Account Settings, Provider updates it in memory and immediately reflects the new name across Home, Profile, and Settings without restarting the app or passing variables back and forth."),
        
        ("lib/services/", "Outside World Communication Layer", 
         "Handles background jobs: Firebase Authentication, Cloud Firestore database syncing, HTTP network calls, saving local settings, and launching external phone/browser links. "
         "Why: Screens should only worry about UI. If we ever change our database or API, we only edit this folder without touching the screens."),
        
        ("lib/data/", "Catalog & Mock Registries", 
         "Stores large data sets (like 14 top categories, 2,000+ brand catalogs, banners, and offline deals). "
         "Why: Keeps huge lists of store data cleanly separated so they don't clutter UI screen files."),
        
        ("lib/main.dart", "The Front Door of the App", 
         "The starting point of the application where themes, state providers, and all 30 screen route names are registered.")
    ]

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Folder"
    hdr[1].text = "Role"
    hdr[2].text = "Why We Use It (Reasoning)"
    for c in hdr:
        set_cell_background(c, "1E90FF")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        c.paragraphs[0].runs[0].font.bold = True

    for f_name, f_role, f_why in folders:
        row = table.add_row().cells
        row[0].text = f_name
        row[1].text = f_role
        row[2].text = f_why
        set_cell_background(row[0], "F1F5F9")
        set_cell_padding(row[0], 80, 80, 100, 100)
        set_cell_padding(row[1], 80, 80, 100, 100)
        set_cell_padding(row[2], 80, 80, 100, 100)

    doc.add_paragraph()

    # SECTION 3: COMPLETE APP WORKFLOW
    add_heading(doc, "3. Complete App Workflow (Step-by-Step)", level=1)
    
    workflow_steps = [
        ("Step 1: App Launch & Authentication", 
         "User opens the app -> Splash Screen checks if session exists. If new, Onboarding Screen explains benefits, then Login Screen authenticates via Phone OTP or Email. Profile data syncs with Cloud Firestore."),
        
        ("Step 2: Browsing & Discovering Deals", 
         "User lands on Home Screen -> Can search items, scroll 14 top categories (Fashion, Mobiles, Pharmacy, Loans, etc.), check flash deal reels (Flipkart, Meesho Under ₹99), or scratch the Golden Ticket for bonus cash."),
        
        ("Step 3: Visiting Retailer & Tracking Order", 
         "User selects a product or store -> Views effective discounted price on Product Detail Screen -> Taps 'Shop Now' -> Shopping Confirmation screen activates tracking session -> Redirects user into the retailer's app/website (Amazon, Flipkart, etc.) where they complete the purchase."),
        
        ("Step 4: Cashback Lifecycle (Pending to Confirmed)", 
         "Retailer reports the purchase within 24-48 hours -> Cashback appears in user's Pending Earnings. After retailer's 30-day return period ends, the cashback automatically shifts into Confirmed Earnings."),
        
        ("Step 5: Withdrawing Real Money", 
         "User opens My Earnings -> Taps 'Request Payment' -> Selects UPI ID, Bank NEFT, or Amazon Pay Gift Card -> Enters amount (minimum ₹250) -> Confirms withdrawal."),
        
        ("Step 6: Referral Bonus & Support", 
         "User shares their unique invite link from Refer & Earn -> Friends join and shop -> User earns a 10% lifetime commission. If any purchase wasn't tracked, user raises a claim in Missing Cashback Tickets.")
    ]

    for s_title, s_desc in workflow_steps:
        p = doc.add_paragraph()
        r1 = p.add_run(f"• {s_title}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(30, 144, 255)
        r2 = p.add_run(s_desc)
        r2.font.size = Pt(10.5)

    doc.add_paragraph()

    # SECTION 4: ALL 30 SCREENS CONCISE BREAKDOWN
    add_heading(doc, "4. Complete Breakdown of All 30 Screens", level=1)

    screen_groups = [
        ("Group 1: Onboarding & Authentication", [
            ("1. Splash Screen (/splash)", "Entry screen; displays pulsating logo and checks if user is logged in to route to Onboarding or Home."),
            ("2. Onboarding Screen (/onboarding)", "3-page introductory swipe slider explaining Cashback, Instant Withdrawals, and 10% Referrals."),
            ("3. Login & Sign Up Screen (/login)", "Dual-mode login supporting Phone Number OTP verification and Email/Password login, plus registration bottom sheet."),
            ("4. Sign Up Screen (/signup)", "Lightweight direct route alias forwarding new users into the registration flow.")
        ]),
        ("Group 2: Home & Shopping Marketplace", [
            ("5. Home Screen (/ or /home)", "Main app hub with live search bar, 14 categories, Golden Ticket banner, store deal reels, and 12 discovery sections."),
            ("6. Live Search Screen (/search)", "Instant real-time search filtering across 2,000+ brand deals, products, and categories as you type."),
            ("7. Categories Screen (/categories)", "Product catalog with a top auto-centering category bar and a live feed of discounted products."),
            ("8. All Categories Screen (/all-categories)", "Visual 2-column image grid showing all 24 eCommerce shopping categories with deal counters."),
            ("9. Top Category Brands Screen (/top-category-brands)", "Lists all partner stores for a specific category (e.g. Fashion or Pharmacy) with exit confirmation dialogs."),
            ("10. Offer Section Deals Screen (/offer-section)", "Grid displaying special promotional sales (Flipkart Deals, Meesho Under ₹99, Best of Loans)."),
            ("11. Product Detail Screen (/product-detail)", "Full offer page with image carousel, effective price math ('Final Price = Price - Cashback'), 3-step guide, and 'Shop Now' button."),
            ("12. Shopping Confirmation Screen (/shopping-confirmation)", "Transition screen confirming cashback tracking before opening the retailer website/app."),
            ("13. Golden Ticket Screen (/ticket)", "Interactive gamified golden ticket with real-time 3D wave physics and gesture scratch reveal for vouchers.")
        ]),
        ("Group 3: Profile, Wallet & Payouts", [
            ("14. Profile Dashboard Screen (/profile)", "Account control center showing user info, total cashback stats, menu links, dark mode switch, and logout."),
            ("15. Account Settings Screen (/account-settings)", "Form allowing users to edit and update their Full Name, Email, and Phone number with Cloud Firestore sync."),
            ("16. My Earnings Screen (/my-earnings)", "Financial wallet breakdown showing Lifetime Earnings, Confirmed Balance (ready to withdraw), and Pending Balance."),
            ("17. Why Is Earnings Pending Screen (/know-why)", "Educational 4-step timeline explaining tracking, return period wait, store validation, and confirmation."),
            ("18. Withdraw Screen (/withdraw)", "Payout interface allowing users to transfer confirmed money (min ₹250) to UPI, Bank Account, or Amazon Pay."),
            ("19. Payments Guide Screen (/payments)", "Informational guide explaining withdrawal thresholds, payment rules, and payout methods."),
            ("20. Payments History Screen (/payments-history)", "Searchable statement of all past payouts with a Digital Receipt popup showing bank reference IDs."),
            ("21. My Order Details Screen (/my-order-details)", "Transaction history ledger of all tracked store purchases with status, dates, and cashback earned.")
        ]),
        ("Group 4: Referrals, Customer Support & Legal", [
            ("22. Refer & Earn Screen (/refer-earn)", "Referral hub with auto-sliding banner, 1-tap link copy button, WhatsApp sharing, and 10% lifetime commission guide."),
            ("23. My Referrals Screen (/my-referrals)", "Tracker showing total invited friends who joined, total referral bonus earned, and list of referrals."),
            ("24. Missing Cashback Tickets Screen (/missing-tickets)", "Claim tool for untracked store orders; allows submitting retailer name, order ID, and amount for resolution."),
            ("25. Get Help & Support Screen (/get-help)", "Customer FAQ center with categorized topics (Shopping, Tracking, Payouts) and contact support modal."),
            ("26. Call Us Helpline Screen (/call-us)", "Direct phone contact portal with 1-tap toll-free helpline caller (1800-227-4527), hours, email, and WhatsApp."),
            ("27. Your Queries Screen (/your-queries)", "Overview of open support queries and missing ticket claims with quick navigation links."),
            ("28. Review Us Screen (/review-us)", "5-star rating picker, feedback submission box, Play Store link, and community review feed with upvote buttons."),
            ("29. Notifications Center Screen (/notifications)", "Filterable alert inbox for cashback updates, flash sales, and payouts with 'Mark all as read'."),
            ("30. Privacy Policy Screen (/privacy-policy)", "10-section structured legal policy covering data protection, encryption, account security, and user rights.")
        ])
    ]

    for g_title, g_screens in screen_groups:
        add_heading(doc, g_title, level=2)
        for s_title, s_desc in g_screens:
            p = doc.add_paragraph()
            r1 = p.add_run(f"• {s_title}: ")
            r1.bold = True
            r1.font.color.rgb = RGBColor(30, 144, 255)
            r2 = p.add_run(s_desc)
            r2.font.size = Pt(10)

    doc.add_paragraph()

    # SECTION 5: CATEGORIES & LOANS OVERVIEW
    add_heading(doc, "5. Categories & Loans Vertical Overview", level=1)
    doc.add_paragraph(
        "• 14 Featured Home Categories: Most Popular, Fashion, Credit Cards, Beauty & Grooming, Home & Kitchen, "
        "Electronics, Food & Grocery, Mobiles, Pharmacy, Health & Wellness, Loans & Financial Services, Departmental, "
        "Flights & Hotels, Education."
    )
    doc.add_paragraph(
        "• Loans & Financial Services: Features Instant Personal Loans (Navi, MoneyView, KreditBee), Pre-Approved Home Loans, "
        "and Zero-Interest Credit Cards. When a user's loan application is approved by the partner lender, flat cash rewards are "
        "automatically credited to the user's wallet."
    )

    doc.save(DOC_PATH)
    print(f"[SUCCESS] Clean concise docx created at: {DOC_PATH}")

if __name__ == "__main__":
    build_word_doc()
